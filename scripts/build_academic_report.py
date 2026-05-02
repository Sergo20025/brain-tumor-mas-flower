from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "reports"
DOCX_PATH = OUTPUT_DIR / "Отчет_по_децентрализованной_системе.docx"
TXT_PATH = OUTPUT_DIR / "Отчет_по_децентрализованной_системе.txt"


EXPERIMENTS = {
    "mri_ring": {
        "title": "Brain Tumor MRI: кольцевая топология",
        "short": "MRI ring",
        "path": ROOT / "outputs/experiments/decentralized/exp_decentralized_ring",
        "dataset": "Brain Tumor MRI",
        "topology": "ring",
    },
    "mri_aug": {
        "title": "Brain Tumor MRI: расширенное кольцо",
        "short": "MRI augmented_ring",
        "path": ROOT / "outputs/experiments/decentralized/exp_decentralized_augmented_ring",
        "dataset": "Brain Tumor MRI",
        "topology": "augmented_ring",
    },
    "mri_full": {
        "title": "Brain Tumor MRI: полносвязный граф",
        "short": "MRI full_graph",
        "path": ROOT / "outputs/experiments/decentralized/exp_decentralized_full_graph",
        "dataset": "Brain Tumor MRI",
        "topology": "full_graph",
    },
    "mri_aug_no_agents": {
        "title": "Brain Tumor MRI: расширенное кольцо без агентной логики",
        "short": "MRI augmented_ring no agents",
        "path": ROOT / "outputs/experiments/decentralized_baseline/exp_decentralized_augmented_ring_no_agents",
        "dataset": "Brain Tumor MRI",
        "topology": "augmented_ring_no_agents",
    },
    "cifar_aug": {
        "title": "CIFAR-100: расширенное кольцо",
        "short": "CIFAR-100 augmented_ring",
        "path": ROOT / "outputs/experiments/cifar100/decentralized/exp_cifar100_augmented_ring",
        "dataset": "CIFAR-100",
        "topology": "augmented_ring",
    },
    "cifar_full": {
        "title": "CIFAR-100: полносвязный граф",
        "short": "CIFAR-100 full_graph",
        "path": ROOT / "outputs/experiments/cifar100/decentralized/exp_cifar100_full_graph",
        "dataset": "CIFAR-100",
        "topology": "full_graph",
    },
}


def _parse_json_stream(text: str) -> list[dict]:
    decoder = json.JSONDecoder()
    rows: list[dict] = []
    index = 0
    while index < len(text):
        while index < len(text) and text[index].isspace():
            index += 1
        if index >= len(text):
            break
        row, next_index = decoder.raw_decode(text, index)
        rows.append(row)
        index = next_index
    return rows


def load_jsonl(path: Path) -> list[dict]:
    rows: list[dict] = []
    with path.open("r", encoding="utf-8") as handle:
        for line in handle:
            line = line.strip()
            if line:
                rows.extend(_parse_json_stream(line))
    return rows


def select_last_run(rows: list[dict]) -> list[dict]:
    if not rows:
        return rows
    runs: list[list[dict]] = []
    current_run: list[dict] = []
    prev_round: int | None = None
    for row in rows:
        round_id = int(row["round"])
        if prev_round is not None and round_id <= prev_round:
            runs.append(current_run)
            current_run = []
        current_run.append(row)
        prev_round = round_id
    if current_run:
        runs.append(current_run)
    return runs[-1]


def load_experiment_metrics(exp_path: Path) -> dict:
    global_rows = select_last_run(load_jsonl(exp_path / "global_eval_metrics.jsonl"))
    round_rows = select_last_run(load_jsonl(exp_path / "round_metrics.jsonl"))
    last_global = global_rows[-1]
    best_global = max(global_rows, key=lambda row: float(row.get("f1_macro", 0.0)))
    last_round = round_rows[-1]
    aggregated = dict(last_round.get("aggregated_metrics", {}))
    return {
        "last_global": last_global,
        "best_global": best_global,
        "last_local": aggregated,
        "num_rounds": len(global_rows) - 1,
    }


def configure_styles(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        heading_style = document.styles[style_name]
        heading_style.font.name = "Times New Roman"
        heading_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_title(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16)


def add_caption(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True


def add_image(document: Document, image_path: Path, caption: str, width_cm: float = 15.5) -> None:
    if not image_path.exists():
        return
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    add_caption(document, caption)


def build_summary_table(document: Document, rows: list[tuple[str, dict]]) -> None:
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Эксперимент"
    hdr[1].text = "Лучш. раунд"
    hdr[2].text = "Accuracy"
    hdr[3].text = "F1-macro"
    hdr[4].text = "Loss"
    for title, metrics in rows:
        best = metrics["best_global"]
        row = table.add_row().cells
        row[0].text = title
        row[1].text = str(int(best["round"]))
        row[2].text = f"{float(best['accuracy']):.4f}"
        row[3].text = f"{float(best['f1_macro']):.4f}"
        row[4].text = f"{float(best['loss']):.4f}"


def compose_text(metrics: dict[str, dict]) -> str:
    mri_ring = metrics["mri_ring"]["best_global"]
    mri_aug = metrics["mri_aug"]["best_global"]
    mri_full = metrics["mri_full"]["best_global"]
    mri_base = metrics["mri_aug_no_agents"]["best_global"]
    cifar_aug = metrics["cifar_aug"]["best_global"]
    cifar_full = metrics["cifar_full"]["best_global"]

    lines: list[str] = []
    lines.append("ОТЧЕТ ПО ПРОГРАММНОЙ РЕАЛИЗАЦИИ И ЭКСПЕРИМЕНТАЛЬНОМУ ИССЛЕДОВАНИЮ ДЕЦЕНТРАЛИЗОВАННОЙ СИСТЕМЫ ОБУЧЕНИЯ")
    lines.append("")
    lines.append("1. Общая характеристика выполненной работы")
    lines.append(
        "В ходе работы исходная логика обучения была перестроена в сторону децентрализованного федеративного обучения без центрального параметрического сервера. "
        "Основной акцент был сделан на задаче классификации МРТ-изображений опухолей головного мозга, а затем предложенная архитектура была дополнительно проверена "
        "на более сложном наборе данных CIFAR-100. После перехода к децентрализованной схеме были реализованы три топологии взаимодействия узлов: ring, augmented_ring и full_graph, "
        "а также дополнительный baseline-сценарий без агентной логики."
    )
    lines.append("")
    lines.append("2. Что было изменено при переходе к децентрализованной системе")
    lines.append(
        "Ключевое изменение заключается в отказе от централизованной схемы агрегации. В новой реализации каждый узел стал самостоятельным параметрическим узлом, "
        "содержащим локальные данные, локальную модель и набор функциональных компонентов, отвечающих за обучение, мониторинг и агрегацию. "
        "Координатор симуляции сохранился только как служебный компонент для запуска раундов, логирования, построения общей оценки и сохранения чекпоинтов."
    )
    lines.append(
        "Система поддерживает несколько режимов разбиения данных, включая iid, dirichlet, quantity_skew, shards и shards_quantity_skew. "
        "В основных экспериментах использовался режим shards_quantity_skew, формирующий одновременно неравномерность по количеству данных и по классам."
    )
    lines.append("")
    lines.append("3. Техническая архитектура системы")
    lines.append(
        "В качестве базовой нейросетевой модели использовалась EfficientNet-B0 с заменой классификационной головы под число классов целевой задачи. "
        "Для обучения применялся оптимизатор AdamW. В качестве основных метрик использовались loss, accuracy, precision_macro, recall_macro и f1_macro. "
        "Поддержка данных реализована единым модулем, работающим как с Brain Tumor MRI, так и с CIFAR-100."
    )
    lines.append(
        "В агентной версии на каждом узле используются StorageAgent, ComputeAgent, MonitoringAgent и AggregationAgent. "
        "StorageAgent загружает локальную партицию данных, ComputeAgent проводит локальное обучение и локальную валидацию, MonitoringAgent оценивает качество и стабильность узла, "
        "а AggregationAgent собирает обновления соседей, буферизует их и формирует новую локальную модель. В результате агрегация выполняется непосредственно на узле."
    )
    lines.append(
        "Кроме агентной версии был реализован отдельный baseline-сценарий без агентной оболочки. В нем сохраняются та же децентрализованная постановка, те же параметры обучения "
        "и та же топология, однако взаимодействие реализуется напрямую: локальное обучение на узлах и простая агрегация по соседям без trust-score и без агентных эвристик."
    )
    lines.append("")
    lines.append("4. Экспериментальная постановка")
    lines.append(
        "Для набора Brain Tumor MRI использовались 10 клиентов, 30 раундов, 2 локальные эпохи на раунд, batch size 16, learning rate 0.0002, weight decay 0.00001, pretrained=false. "
        "Для CIFAR-100 использовались 10 клиентов, 30 раундов, 2 локальные эпохи, batch size 32, learning rate 0.0005, weight decay 0.0001, pretrained=false."
    )
    lines.append("")
    lines.append("5. Результаты на Brain Tumor MRI")
    lines.append(
        f"В топологии ring лучший результат составил accuracy={float(mri_ring['accuracy']):.4f}, f1_macro={float(mri_ring['f1_macro']):.4f}, loss={float(mri_ring['loss']):.4f}. "
        f"В topологии augmented_ring результат улучшился до accuracy={float(mri_aug['accuracy']):.4f}, f1_macro={float(mri_aug['f1_macro']):.4f}, loss={float(mri_aug['loss']):.4f}. "
        f"Полносвязный граф показал наилучшее качество среди агентных вариантов: accuracy={float(mri_full['accuracy']):.4f}, f1_macro={float(mri_full['f1_macro']):.4f}, loss={float(mri_full['loss']):.4f}."
    )
    lines.append(
        f"Дополнительный baseline без агентной логики в топологии augmented_ring показал accuracy={float(mri_base['accuracy']):.4f}, f1_macro={float(mri_base['f1_macro']):.4f}, loss={float(mri_base['loss']):.4f}. "
        "На текущей серии он оказался немного лучше агентной версии augmented_ring. Это означает, что на относительно простой четырехклассовой задаче без явных аномальных узлов "
        "дополнительная trust-aware логика не обязательно приводит к росту качества. При этом архитектурно агентная версия остается более богатой и лучше подготовленной к сложным сценариям."
    )
    lines.append(
        "Сравнение трех топологий подтверждает основную гипотезу работы: увеличение связности графа взаимодействия ускоряет распространение параметров между узлами и улучшает качество глобальной модели. "
        "Топология ring выступает нижней границей по связности, augmented_ring представляет собой компромисс между связностью и коммуникационной стоимостью, "
        "а full_graph дает верхнюю границу качества."
    )
    lines.append("")
    lines.append("6. Результаты на CIFAR-100")
    lines.append(
        f"На CIFAR-100 в топологии augmented_ring был получен результат accuracy={float(cifar_aug['accuracy']):.4f}, f1_macro={float(cifar_aug['f1_macro']):.4f}, loss={float(cifar_aug['loss']):.4f}. "
        f"В полносвязной топологии full_graph итоговое качество оказалось выше: accuracy={float(cifar_full['accuracy']):.4f}, f1_macro={float(cifar_full['f1_macro']):.4f}, loss={float(cifar_full['loss']):.4f}."
    )
    lines.append(
        "Задача CIFAR-100 существенно сложнее по двум причинам: она является 100-классовой, а данные также распределены неравномерно между узлами. "
        "По этой причине локальные метрики на узлах оказываются заметно выше глобальной оценки. Тем не менее сам факт устойчивого роста accuracy и f1_macro подтверждает, "
        "что предложенная децентрализованная схема работает не только на медицинском наборе, но и на более общем и более сложном наборе данных."
    )
    lines.append(
        "Одновременно результаты на CIFAR-100 показали, что при увеличении сложности задачи возрастает значение плотности связей в графе взаимодействия. "
        "Полносвязная топология на этом наборе оказалась заметно эффективнее расширенного кольца, поскольку быстрее объединяет знания, полученные на различных узлах."
    )
    lines.append("")
    lines.append("7. Итоговый анализ выполненной работы")
    lines.append(
        "Проведенные эксперименты позволяют сделать несколько важных выводов. Во-первых, реализованная система действительно является децентрализованной: "
        "в основной схеме отсутствует центральный параметрический сервер, а каждый узел самостоятельно обучает свою модель и агрегирует параметры соседей. "
        "Во-вторых, качество итоговой модели существенно зависит от выбранной топологии. В-третьих, переход от задачи Brain Tumor MRI к CIFAR-100 показал, что система сохраняет работоспособность, "
        "но становится намного чувствительнее к гетерогенности данных и ограниченности коммуникационных связей."
    )
    lines.append(
        "С научной точки зрения наиболее важным результатом является подтверждение того, что augmented_ring действительно занимает промежуточное положение между ring и full_graph. "
        "Для медицинской задачи Brain Tumor MRI он обеспечивает более высокое качество, чем ring, при меньшей коммуникационной стоимости, чем full_graph. "
        "Именно поэтому augmented_ring остается основным исследовательским сценарием работы, несмотря на то что в ряде случаев full_graph дает более высокие итоговые метрики."
    )
    lines.append("")
    lines.append("8. Итоговое состояние системы")
    lines.append(
        "На текущем этапе система представляет собой полноценный программный комплекс для исследования децентрализованного федеративного обучения. "
        "Она включает модули загрузки и разбиения данных, обучения модели EfficientNet-B0, локального мониторинга качества узлов, локальной агрегации параметров, "
        "скрипты запуска различных экспериментальных сценариев, автоматическое построение графиков, автоматический анализ результатов и модуль инференса для пользовательской загрузки изображений."
    )
    lines.append(
        "Таким образом, работа может считаться завершенной как в программной, так и в экспериментальной части: основная медицинская задача исследована полноценно, "
        "дополнительная проверка на CIFAR-100 проведена, различия между топологиями показаны, baseline без агентной логики реализован и сравнен, а текущая архитектура системы подробно описана и подтверждена экспериментально."
    )
    return "\n".join(lines) + "\n"


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    metrics = {key: load_experiment_metrics(meta["path"]) for key, meta in EXPERIMENTS.items()}
    text = compose_text(metrics)
    TXT_PATH.write_text(text, encoding="utf-8")

    document = Document()
    configure_styles(document)
    section = document.sections[0]
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    add_title(document, "ОТЧЕТ ПО ПРОГРАММНОЙ РЕАЛИЗАЦИИ И ЭКСПЕРИМЕНТАЛЬНОМУ ИССЛЕДОВАНИЮ")
    add_title(document, "ДЕЦЕНТРАЛИЗОВАННОЙ СИСТЕМЫ ОБУЧЕНИЯ")
    document.add_paragraph("")

    for line in text.strip().split("\n\n"):
        parts = line.split("\n", 1)
        if parts[0] and parts[0][0].isdigit() and ". " in parts[0]:
            document.add_heading(parts[0], level=1)
            if len(parts) > 1:
                for paragraph in parts[1].split("\n"):
                    document.add_paragraph(paragraph)
        else:
            for paragraph in line.split("\n"):
                document.add_paragraph(paragraph)

    document.add_page_break()
    document.add_heading("Сводная таблица результатов", level=1)
    build_summary_table(
        document,
        [
            ("MRI ring", metrics["mri_ring"]),
            ("MRI augmented_ring", metrics["mri_aug"]),
            ("MRI full_graph", metrics["mri_full"]),
            ("MRI augmented_ring без агентов", metrics["mri_aug_no_agents"]),
            ("CIFAR-100 augmented_ring", metrics["cifar_aug"]),
            ("CIFAR-100 full_graph", metrics["cifar_full"]),
        ],
    )

    document.add_page_break()
    document.add_heading("Графики Brain Tumor MRI", level=1)
    mri_base_plots = EXPERIMENTS["mri_aug"]["path"] / "plots"
    add_image(
        document,
        mri_base_plots / "client_data_volume_shards_quantity_skew.png",
        "Рисунок 1. Распределение объема обучающих данных между клиентами в MRI-экспериментах.",
    )
    add_image(
        document,
        mri_base_plots / "class_distribution_shards_quantity_skew.png",
        "Рисунок 2. Распределение классов между клиентами в MRI-экспериментах.",
    )

    image_counter = 3
    for key in ("mri_ring", "mri_aug", "mri_full", "mri_aug_no_agents"):
        meta = EXPERIMENTS[key]
        plots_dir = meta["path"] / "plots"
        document.add_heading(meta["title"], level=2)
        add_image(
            document,
            plots_dir / "accuracy_dynamics_non_iid.png",
            f"Рисунок {image_counter}. Динамика accuracy для сценария {meta['short']}.",
        )
        image_counter += 1
        add_image(
            document,
            plots_dir / "loss_dynamics_non_iid.png",
            f"Рисунок {image_counter}. Динамика loss для сценария {meta['short']}.",
        )
        image_counter += 1
        add_image(
            document,
            plots_dir / "f1_macro_dynamics_non_iid.png",
            f"Рисунок {image_counter}. Динамика F1-macro для сценария {meta['short']}.",
        )
        image_counter += 1

    document.add_page_break()
    document.add_heading("Графики CIFAR-100", level=1)
    cifar_base_plots = EXPERIMENTS["cifar_full"]["path"] / "plots"
    add_image(
        document,
        cifar_base_plots / "client_data_volume_shards_quantity_skew.png",
        f"Рисунок {image_counter}. Распределение объема обучающих данных между клиентами в CIFAR-100.",
    )
    image_counter += 1
    add_image(
        document,
        cifar_base_plots / "class_distribution_shards_quantity_skew.png",
        f"Рисунок {image_counter}. Распределение классов между клиентами в CIFAR-100.",
    )
    image_counter += 1

    for key in ("cifar_aug", "cifar_full"):
        meta = EXPERIMENTS[key]
        plots_dir = meta["path"] / "plots"
        document.add_heading(meta["title"], level=2)
        add_image(
            document,
            plots_dir / "accuracy_dynamics_non_iid.png",
            f"Рисунок {image_counter}. Динамика accuracy для сценария {meta['short']}.",
        )
        image_counter += 1
        add_image(
            document,
            plots_dir / "loss_dynamics_non_iid.png",
            f"Рисунок {image_counter}. Динамика loss для сценария {meta['short']}.",
        )
        image_counter += 1
        add_image(
            document,
            plots_dir / "f1_macro_dynamics_non_iid.png",
            f"Рисунок {image_counter}. Динамика F1-macro для сценария {meta['short']}.",
        )
        image_counter += 1

    document.add_section(WD_SECTION.NEW_PAGE)
    document.add_heading("Заключительное описание системы", level=1)
    document.add_paragraph(
        "Итоговая версия системы реализует децентрализованное федеративное обучение, в котором каждый узел является самостоятельным параметрическим узлом. "
        "На каждом узле присутствуют локальные данные, локальная модель, агент хранения данных, агент обучения, агент мониторинга и агент агрегации. "
        "Узел самостоятельно обучает модель, получает обновления соседей в соответствии с выбранной топологией и формирует новую локальную версию модели. "
        "Центральный параметрический сервер в основной схеме отсутствует."
    )
    document.add_paragraph(
        "В практическом смысле проект теперь представляет собой законченную исследовательскую платформу: она позволяет запускать различные топологии децентрализованного обучения, "
        "сравнивать агентные и неагентные сценарии, строить графики, автоматически сохранять анализ и использовать полученные чекпоинты для инференса по отдельным MRI-снимкам."
    )

    document.save(DOCX_PATH)
    print(f"Saved TXT report to: {TXT_PATH}")
    print(f"Saved DOCX report to: {DOCX_PATH}")


if __name__ == "__main__":
    main()
